import json
import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "mock-test-print"


PART_TITLES = {
    "L1_picture_choice": "听力第一部分：听录音，选择正确图片。（1-5）",
    "L2_answer_choice": "听力第二部分：听问题，选择合适的回答。（6-10）",
    "L3_dialogue_picture_choice": "听力第三部分：听对话，选择正确图片。（11-15）",
    "L4_sentence_question_answer": "听力第四部分：听句子和问题，选择合适的回答。（16-20）",
    "R1_sentence_picture_match": "阅读第一部分：读句子，选择正确图片。（21-25）",
    "R2_question_answer_match": "阅读第二部分：选择合适的回答。（26-30）",
    "R3_fill_blank": "阅读第三部分：选择合适的词填空。（31-35）",
    "R4_reading_comprehension": "阅读第四部分：读短文，选择正确答案。（36-40）",
}


def all_items(test):
    return [item for section in test["sections"] for item in section["items"]]


def resolve_image(src):
    if not src:
        return None
    src = src.replace("\\", "/")
    if src.startswith("../in-class/"):
        return ROOT / "source" / "in-class" / src[len("../in-class/") :]
    return ROOT / "source" / "data-model" / "mock-tests" / src


def set_cell_text(cell, text, size=10, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc, text="", style=None, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    r.bold = bold
    return p


def setup_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    for name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12.5, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc, test):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{test['titleCn']} - 打印教师包")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("学生试卷 + 听力朗读稿 + 正确答案 | 课堂应急打印版")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(90, 90, 90)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("建议：前半部分打印给学生；后半部分仅教师使用。听力题每题朗读两遍，题间停顿约 4-5 秒。")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10)
    r.font.bold = True


def add_student_header(doc):
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    widths = [Inches(1.0), Inches(2.0), Inches(1.0), Inches(2.0)]
    labels = ["姓名", "", "得分", ""]
    for cell, width, label in zip(t.rows[0].cells, widths, labels):
        cell.width = width
        set_cell_text(cell, label, size=10, bold=bool(label))


def grouped_by_part(items):
    groups = {}
    for item in items:
        groups.setdefault(item["part"], []).append(item)
    return groups


def add_options_table(doc, item, include_images):
    opts = item.get("options", [])
    table = doc.add_table(rows=2 if include_images else 1, cols=len(opts))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for col, opt in enumerate(opts):
        table.cell(0, col).width = Inches(2.1)
        if include_images:
            img = resolve_image(opt.get("image"))
            p = table.cell(0, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img and img.exists():
                run = p.add_run()
                run.add_picture(str(img), width=Inches(1.25))
            else:
                p.add_run("[图片]").font.name = "Microsoft YaHei"
            set_cell_text(table.cell(1, col), opt["id"], size=10, bold=True)
        else:
            set_cell_text(table.cell(0, col), f"{opt['id']}. {opt.get('text','')}", size=10)


def add_image_option_grid(doc, options):
    cols = 3
    rows = ((len(options) + cols - 1) // cols) * 2
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, opt in enumerate(options):
        row = (idx // cols) * 2
        col = idx % cols
        table.cell(row, col).width = Inches(2.1)
        img = resolve_image(opt.get("image"))
        p = table.cell(row, col).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img and img.exists():
            p.add_run().add_picture(str(img), width=Inches(1.15))
        else:
            p.add_run("[图片]").font.name = "Microsoft YaHei"
        set_cell_text(table.cell(row + 1, col), opt["id"], size=10, bold=True)


def add_text_option_row(doc, options):
    table = doc.add_table(rows=1, cols=len(options))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for col, opt in enumerate(options):
        set_cell_text(table.cell(0, col), f"{opt['id']}. {opt.get('text', '')}", size=10, bold=True)


def add_answer_line(doc, item, show_question=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f"{item['number']}. ")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10.5)
    if show_question and item.get("question"):
        r = p.add_run(item["question"])
        r.font.name = "Microsoft YaHei"
        r.font.size = Pt(10.5)
    r = p.add_run("    答案：______")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10.5)


def add_group_picture_part(doc, items, show_questions):
    add_image_option_grid(doc, items[0].get("options", []))
    for item in items:
        add_answer_line(doc, item, show_question=show_questions)


def add_question_answer_match_part(doc, items):
    options = items[0].get("options", [])
    table = doc.add_table(rows=max(len(items), len(options)), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, row in enumerate(table.rows):
        left = ""
        right = ""
        if idx < len(items):
            item = items[idx]
            left = f"{item['number']}. {item.get('question', '')}    ______"
        if idx < len(options):
            opt = options[idx]
            right = f"{opt['id']}. {opt.get('text', '')}"
        set_cell_text(row.cells[0], left, size=10, bold=False)
        set_cell_text(row.cells[1], right, size=10, bold=False)


def add_fill_blank_part(doc, items):
    add_text_option_row(doc, items[0].get("options", []))
    for item in items:
        add_answer_line(doc, item, show_question=True)


def add_student_item(doc, item):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(f"{item['number']}. ")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10.5)
    question = item.get("question") or ""
    if question:
        r = p.add_run(question)
        r.font.name = "Microsoft YaHei"
        r.font.size = Pt(10.5)
    include_images = any(opt.get("image") for opt in item.get("options", []))
    add_options_table(doc, item, include_images=include_images)


def add_student_paper(doc, test):
    doc.add_heading("学生试卷 Student Paper", level=1)
    add_student_header(doc)
    add_paragraph(doc, f"时间：{test.get('durationMinutes', 35)} 分钟    共 40 题，每题 1 分。", bold=True)
    add_paragraph(doc, "听力部分由老师朗读。请听题后选择 A、B、C。")

    groups = grouped_by_part(all_items(test))
    order = [
        "L1_picture_choice",
        "L2_answer_choice",
        "L3_dialogue_picture_choice",
        "L4_sentence_question_answer",
        "R1_sentence_picture_match",
        "R2_question_answer_match",
        "R3_fill_blank",
        "R4_reading_comprehension",
    ]
    for part in order:
        if part not in groups:
            continue
        doc.add_heading(PART_TITLES[part], level=2)
        if part.startswith("L"):
            add_paragraph(doc, "老师朗读本部分题目。学生只看选项作答。", size=9.5, color=(90, 90, 90))
        if part == "L3_dialogue_picture_choice":
            add_group_picture_part(doc, groups[part], show_questions=False)
            continue
        if part == "R1_sentence_picture_match":
            add_group_picture_part(doc, groups[part], show_questions=True)
            continue
        if part == "R2_question_answer_match":
            add_question_answer_match_part(doc, groups[part])
            continue
        if part == "R3_fill_blank":
            add_fill_blank_part(doc, groups[part])
            continue
        passage_shown = set()
        for item in groups[part]:
            if item.get("passage") and item["passage"] not in passage_shown:
                passage_shown.add(item["passage"])
                add_paragraph(doc, item["passage"], bold=True, size=11)
            add_student_item(doc, item)


def listening_lines(test):
    lines = []
    items = [i for i in all_items(test) if i["skill"] == "listening"]
    current = None
    for item in items:
        if item["part"] != current:
            current = item["part"]
            lines.append("")
            lines.append(PART_TITLES[current])
        text = item.get("audioText", "")
        lines.append(f"第 {item['number']} 题：{text}")
        lines.append(f"重复：{text}")
        lines.append("（停顿 4-5 秒）")
    return [line for line in lines if line != ""]


def add_listening_script(doc, test):
    doc.add_page_break()
    doc.add_heading("教师听力朗读稿 Listening Script", level=1)
    add_paragraph(doc, "使用方法：每题读两遍；两遍之间可停 1-2 秒；每题结束后停 4-5 秒。", bold=True)
    current = None
    for item in [i for i in all_items(test) if i["skill"] == "listening"]:
        if item["part"] != current:
            current = item["part"]
            doc.add_heading(PART_TITLES[current], level=2)
        add_paragraph(doc, f"第 {item['number']} 题", bold=True, color=(31, 77, 120))
        add_paragraph(doc, item.get("audioText", ""), size=12)
        add_paragraph(doc, f"重复：{item.get('audioText', '')}", size=11, color=(90, 90, 90))
        add_paragraph(doc, "停顿 4-5 秒", size=9.5, color=(120, 120, 120))


def add_answer_key(doc, test):
    doc.add_page_break()
    doc.add_heading("正确答案 Answer Key", level=1)
    rows = [all_items(test)[i : i + 10] for i in range(0, len(all_items(test)), 10)]
    table = doc.add_table(rows=1, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_text(cell, str(i + 1), size=9, bold=True)
    for chunk in rows:
        nums = table.add_row().cells
        ans = table.add_row().cells
        for idx, item in enumerate(chunk):
            set_cell_text(nums[idx], str(item["number"]), size=9, bold=True)
            set_cell_text(ans[idx], item["answer"], size=11, bold=True)

    doc.add_heading("分题答案", level=2)
    for item in all_items(test):
        text = f"{item['number']}. {item['answer']}"
        if item.get("question"):
            text += f"    {item['question']}"
        elif item.get("audioText"):
            text += f"    {item['audioText']}"
        add_paragraph(doc, text, size=10)


def write_text_files(test, script_path, answer_path):
    script = [f"{test.get('titleCn', test['title'])} - 听力朗读稿", "每题读两遍；题间停顿 4-5 秒。", ""]
    script.extend(listening_lines(test))
    script_path.write_text("\n".join(script) + "\n", encoding="utf-8")

    answers = [f"{test.get('titleCn', test['title'])} - 正确答案", ""]
    answers.append(" ".join(f"{item['number']}.{item['answer']}" for item in all_items(test)))
    answers.append("")
    for item in all_items(test):
        answers.append(f"{item['number']}. {item['answer']}")
    answer_path.write_text("\n".join(answers) + "\n", encoding="utf-8")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--test-id", default="HSK1-mock-01")
    args = parser.parse_args()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data_path = ROOT / "source" / "data-model" / "mock-tests" / f"{args.test_id}.json"
    docx_path = OUT_DIR / f"{args.test_id}-print-teacher-pack.docx"
    script_path = OUT_DIR / f"{args.test_id}-listening-script.txt"
    answer_path = OUT_DIR / f"{args.test_id}-answer-key.txt"

    test = json.loads(data_path.read_text(encoding="utf-8"))
    doc = Document()
    setup_document(doc)
    add_title(doc, test)
    add_student_paper(doc, test)
    add_listening_script(doc, test)
    add_answer_key(doc, test)
    doc.save(docx_path)
    write_text_files(test, script_path, answer_path)
    print(docx_path)
    print(script_path)
    print(answer_path)


if __name__ == "__main__":
    main()
